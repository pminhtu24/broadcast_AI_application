import os
from datetime import datetime
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap


def number_to_vietnamese_words(num: float) -> str:
    if num == 0:
        return "Không đồng"

    units = ["", "nghìn", "triệu", "tỷ"]
    result = ""

    num_str = str(int(num))
    length = len(num_str)

    if length <= 3:
        result = str(int(num))
        return result

    groups = []
    while num_str:
        groups.append(num_str[-3:])
        num_str = num_str[:-3]

    groups.reverse()

    for i, group in enumerate(groups):
        if int(group) != 0:
            result += f"{int(group)} {units[len(groups) - i - 1]} "

    result = result.strip()
    result = result.replace("nghìn", "nghìn")
    result = result.replace("triệu", "triệu")
    result = result.replace("tỷ", "tỷ")

    return result


def set_cell_shading(cell, color: str):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_font_for_paragraph(paragraph, font_name: str = "Times New Roman", font_size: int = 12):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def set_document_font(doc: Document, font_name: str = "Times New Roman"):
    """Set default font for entire document"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
        if not paragraph.runs and paragraph.text:
            paragraph.clear()
            run = paragraph.add_run(paragraph.text)
            run.font.name = font_name
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name


def format_currency(amount: float) -> str:
    return f"{amount:,.0f}".replace(",", ".")


class QuoteGenerator:
    def __init__(self, output_dir: str = "backend/generated_quotes"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(
        self,
        customer_name: str,
        items: List[Dict[str, Any]],
        price_list: str,
        customer_address: Optional[str] = None,
        contract_start: Optional[str] = None,
        contract_end: Optional[str] = None,
        session_id: Optional[str] = None
    ) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"bao_gia_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)

        doc = Document()
        set_document_font(doc, "Times New Roman")

        self._set_page_margins(doc)
        self._add_header(doc)
        self._add_title(doc, price_list)
        self._add_date(doc)
        self._add_recipient(doc, customer_name, customer_address, price_list)
        self._add_contract_info(doc, contract_start, contract_end)
        self._add_service_intro(doc, price_list)
        self._add_table(doc, items, price_list)
        self._add_total_and_words(doc, items, price_list)
        self._add_contact(doc)
        self._add_signature(doc)

        doc.save(filepath)
        return filename

    def _set_page_margins(self, doc: Document):
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _add_header(self, doc: Document):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell_left = table.cell(0, 0)
        cell_right = table.cell(0, 1)

        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_left.add_run("THÀNH ỦY HẢI PHÒNG\n")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run = p_left.add_run("BÁO VÀ PHÁT THANH, TRUYỀN HÌNH HẢI PHÒNG")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_right.add_run("ĐẢNG CỘNG SẢN VIỆT NAM\n")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run = p_right.add_run("Độc lập - Tự do - Hạnh phúc")
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    def _add_title(self, doc: Document, price_list: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if price_list == "413":
            run = p.add_run("BÁO GIÁ HỖ TRỢ TUYÊN TRUYỀN")
        else:
            run = p.add_run("BÁO GIÁ QUẢNG CÁO")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

    def _add_date(self, doc: Document):
        now = datetime.now()
        p = doc.add_paragraph()
        p.add_run(f"Ngày {now.day:02d} tháng {now.month:02d} năm {now.year}")

    def _add_recipient(self, doc: Document, customer_name: str, customer_address: Optional[str], price_list: str):
        p = doc.add_paragraph()
        run = p.add_run(f"Kính gửi: {customer_name}")
        run.font.name = "Times New Roman"

        if customer_address:
            p2 = doc.add_paragraph()
            run = p2.add_run(customer_address)
            run.font.name = "Times New Roman"

        if price_list == "413":
            p3 = doc.add_paragraph()
            run = p3.add_run("Lời đầu tiên xin chân thành cảm ơn Quý cơ quan, chúng tôi xin được gửi báo giá hỗ trợ tuyên truyền như sau:")
            run.font.name = "Times New Roman"

    def _add_contract_info(self, doc: Document, contract_start: Optional[str], contract_end: Optional[str]):
        p = doc.add_paragraph()
        content = "- Thời gian thực hiện hợp đồng: "

        if contract_start and contract_end:
            content += f"từ ngày {contract_start} đến ngày {contract_end}"
        elif contract_start:
            content += f"từ ngày {contract_start} đến ngày.................."
        elif contract_end:
            content += f"từ ngày.................. đến ngày {contract_end}"
        else:
            content += "từ ngày.................. đến ngày.................."

        run = p.add_run(content)
        run.font.name = "Times New Roman"

        if len(doc.paragraphs) > 0:
            last_p = doc.paragraphs[-1]
            if last_p != p:
                pass
            else:
                p2 = doc.add_paragraph()
                run = p2.add_run("- Thời lượng: 60 giây/1 lần phát sóng")
                run.font.name = "Times New Roman"

    def _add_service_intro(self, doc: Document, price_list: str):
        if price_list == "413":
            p = doc.add_paragraph()
            run = p.add_run("Chúng tôi xin được gửi báo giá hỗ trợ tuyên truyền như sau:")
            run.font.name = "Times New Roman"
        else:
            p = doc.add_paragraph()
            run = p.add_run("Chúng tôi xin được gửi báo giá chi tiết quảng cáo trên kênh truyền hình THP của Báo và phát thanh, truyền hình Hải Phòng như sau:")
            run.font.name = "Times New Roman"

    def _add_table(self, doc: Document, items: List[Dict[str, Any]], price_list: str):
        if price_list == "413":
            headers = ["NỘI DUNG", "THỂ LOẠI", "THỜI LƯỢNG", "SỐ LƯỢNG", "ĐƠN GIÁ"]
            cols = 5
        else:
            headers = ["MỤC", "Khung giờ", "Số lần phát sóng", "Đơn giá", "Thành tiền"]
            cols = 5

        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            set_cell_shading(cell, "D9D9D9")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"

        total = 0
        for item in items:
            row = table.add_row()

            if price_list == "413":
                qty = item.get("quantity", 1)
                unit = item.get("unit_price", 0)
                item_total = qty * unit
                total += item_total

                row.cells[0].text = item.get("content", "")
                row.cells[1].text = item.get("type", "")
                row.cells[2].text = item.get("duration", "")
                row.cells[3].text = str(qty)
                row.cells[4].text = format_currency(item_total)
            else:
                row.cells[0].text = item.get("service", item.get("name", ""))
                row.cells[1].text = item.get("frame", "")
                row.cells[2].text = str(item.get("quantity", ""))
                row.cells[3].text = format_currency(item.get("unit_price", 0))
                row.cells[4].text = format_currency(item.get("total", 0))
                total += item.get("total", 0)

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"

        total_row = table.add_row()
        if price_list == "413":
            total_row.cells[0].merge(total_row.cells[3])
            total_row.cells[0].text = "TỔNG CỘNG"
            total_row.cells[4].text = format_currency(total)
        else:
            total_row.cells[0].merge(total_row.cells[3])
            total_row.cells[0].text = "TỔNG CỘNG"
            total_row.cells[4].text = format_currency(total)

        for cell in total_row.cells:
            set_cell_shading(cell, "E6E6E6")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"

    def _add_total_and_words(self, doc: Document, items: List[Dict[str, Any]], price_list: str):
        total = 0
        for item in items:
            if price_list == "413":
                total += item.get("quantity", 1) * item.get("unit_price", 0)
            else:
                total += item.get("total", 0)
        in_words = number_to_vietnamese_words(total)

        p = doc.add_paragraph()
        run = p.add_run(f"(Bằng chữ: {in_words} đồng)")
        run.font.name = "Times New Roman"

        if price_list != "413":
            p2 = doc.add_paragraph()
            run = p2.add_run("- Giá trên đã bao gồm VAT.")
            run.font.name = "Times New Roman"

    def _add_contact(self, doc: Document):
        p = doc.add_paragraph()
        run = p.add_run("Xin vui lòng liên hệ: Phòng Dịch vụ")
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        run = p.add_run("Địa chỉ: Số 2 Nguyễn Bình, phường Lê Chân, Hải Phòng")
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        run = p.add_run("Điện thoại: 0912.023.771")
        run.font.name = "Times New Roman"

    def _add_signature(self, doc: Document):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("KT.GIÁM ĐỐC")
        run.font.name = "Times New Roman"
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("PHÓ GIÁM ĐỐC")
        run.font.name = "Times New Roman"
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("(Ký và ghi rõ họ tên)")
        run.font.name = "Times New Roman"


def generate_quote_docx(
    customer_name: str,
    items: List[Dict[str, Any]],
    price_list: str,
    customer_address: Optional[str] = None,
    contract_start: Optional[str] = None,
    contract_end: Optional[str] = None
) -> str:
    generator = QuoteGenerator()
    return generator.generate(
        customer_name=customer_name,
        items=items,
        price_list=price_list,
        customer_address=customer_address,
        contract_start=contract_start,
        contract_end=contract_end
    )